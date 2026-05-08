from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
SCREENSHOT_DIR = ROOT / "output" / "screenshots"
OUTPUT_FILE = OUTPUT_DIR / "雾霾探测系统设计实验报告.docx"


def set_document_style(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16 if level == 1 else 13)


def add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_cover_table(document):
    table = document.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    values = [
        ("姓名", "待填写", "学号", "待填写", "学院"),
        ("待填写", "任务分工", "前后端实现、接口联调、测试、报告整理", "贡献度", "100%"),
        ("测试日期", datetime.now().strftime("%Y年%m月%d日"), "指导教师", "待填写", "组号"),
        ("待填写", "电子签名", "提交前插入所有组员签名图片", "", ""),
        ("题目", "雾霾探测系统设计", "", "", ""),
    ]

    for row_index, row_values in enumerate(values):
        for col_index, value in enumerate(row_values):
            table.cell(row_index, col_index).text = value


def add_flow_block(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.3
    run = paragraph.add_run(
        "手机端首页\n"
        "  ↓ 百度地图定位\n"
        "POST /api/location\n"
        "  ↓ 保存定位结果到 app-state.json\n"
        "首页服务端渲染时读取定位并生成 dashboard\n"
        "  ↓ 服务端查询城市 -> 实时天气 -> 空气质量 -> 24 小时预报\n"
        "  ↓ 保存天气快照到 app-state.json\n"
        "首页 Body 区展示天气、AQI 和温湿度折线图"
    )
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(10.5)


def add_screenshot(document, title, filename):
    image_path = SCREENSHOT_DIR / filename
    add_heading(document, title, level=2)
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(5.4))
    else:
        add_body(document, f"截图占位：{filename}。如需最终提交，请在补齐真实 API Key 后重新生成运行截图。")


def add_requirement_table(document):
    add_heading(document, "1.1 题目要求对照", level=2)
    rows = [
        ("任务：手机端雾霾 app 探测系统", "HTML5 手机端单页 Web App，手机浏览器访问后完成定位、天气和空气质量展示。", "已满足"),
        ("要求 1：定位城市保存在服务器端并显示在客户端", "百度地图 JS API 获取城市，POST /api/location 保存到 data/app-state.json，首页 Header 显示定位城市。", "已满足"),
        ("要求 2：天气和空气质量指数动态显示", "首页 Body 直接展示实时天气、AQI、污染物浓度、体感、湿度、风速、能见度。", "已满足"),
        ("要求 3：根据定位城市获取天气详情和 AQI 并保存在服务器端", "服务端读取定位后生成 dashboard，live 模式调用和风天气接口，结果写入 data/app-state.json。", "已满足"),
        ("要求 4：完成报告", "通过 scripts/generate_report.py 生成实验报告 Word 文件。", "已满足，封面信息需手动填写"),
        ("说明 1：可用百度地图 API 获取定位", "前端使用百度地图 JS API，精确定位失败时回退 IP 城市定位，并尝试逆地址解析补省份和城市。", "已满足"),
        ("说明 2：天气和 AQI 可用和风天气等来源", "系统支持 mock 演示和 live 模式；live 模式使用和风天气实时天气、24 小时天气和空气质量 API。", "已满足，真实演示需 API Key"),
        ("说明 3：HTML5 解决手机像素适配，Header 定位，Body 天气/AQI/温湿度折线图", "页面使用 viewport 与响应式 CSS；Header 显示定位，Body 显示天气、AQI、污染物和温湿度折线图。", "已满足"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "原题项目"
    table.rows[0].cells[1].text = "当前实现"
    table.rows[0].cells[2].text = "状态"
    for item, implementation, status in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = implementation
        cells[2].text = status


def add_score_table(document):
    add_heading(document, "5.1 评分点对照", level=2)
    rows = [
        ("定位功能", "20", "城市、省份、经纬度、定位方式保存并回显；服务端持久化定位结果。"),
        ("界面设计", "20", "单页手机天气 App 风格，Header 定位，Body 数据展示，适配手机宽度。"),
        ("天气空气质量指数", "20", "展示实时天气、AQI、污染物浓度、健康提示和温湿度折线图。"),
        ("编码测试", "20", "提供 Flask 路由、数据保存、mock/live 数据服务和自动化测试。"),
        ("方案设计与论证", "10", "报告说明系统结构、接口流程、数据来源和方案取舍。"),
        ("结果与分析", "5", "报告给出页面展示、数据保存和测试结果分析。"),
        ("报告完整性和规范性", "5", "报告按题目格式组织，含问题、方案、数据获取、结果、测试、心得、附录。"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "满分"
    table.rows[0].cells[2].text = "对应实现"
    for item, score, implementation in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = score
        cells[2].text = implementation


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_style(document)

    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    add_title(document, "B级达标测试实验报告")
    add_title(document, "雾霾探测系统设计实验")
    document.add_paragraph("")
    add_cover_table(document)
    document.add_paragraph("")

    add_heading(document, "一、问题描述")
    add_body(
        document,
        "本题要求设计一款手机端雾霾探测系统。系统需要在客户端完成城市定位，并把定位城市保存在服务器端，同时在客户端显示定位结果。服务端根据定位后的城市请求天气详情和空气质量指数，再把结果返回给手机端页面显示，并最终完成实验报告。"
    )
    add_body(
        document,
        "为了严格贴合题目要求，本系统只保留定位、天气详情展示、空气质量指数展示、24小时温湿度折线图以及服务端持久化这几项核心能力，不增加登录、多城市管理、历史记录或地图浏览等题目外功能。"
    )
    add_requirement_table(document)

    add_heading(document, "二、方案设计")
    add_body(
        document,
        "系统采用 Flask 服务端加轻量手机端单页页面的结构。页面 Header 区负责定位城市显示与重新定位，Body 区由 Flask 在服务端组织天气和空气质量数据后渲染。服务端使用本地 JSON 文件保存最近一次定位和天气快照，满足“保存在服务器端”的要求。"
    )
    add_bullet(document, "客户端：同一手机端页面提供定位按钮、天气卡片、AQI 卡片、污染物浓度和温湿度折线图。")
    add_bullet(document, "服务端：使用 Flask 提供 POST /api/location 和 GET /api/dashboard 两个接口。")
    add_bullet(document, "定位能力：百度地图 JSAPI 负责精确定位，失败时退回到 IP 城市定位。")
    add_bullet(document, "天气数据：和风天气 GeoAPI、实时天气、24小时逐小时天气和空气质量 API。")
    add_heading(document, "2.1 系统流程图", level=2)
    add_flow_block(document)

    add_heading(document, "三、数据获取")
    add_body(
        document,
        "定位阶段使用百度地图 JavaScript API 的 Geolocation 能力获取经纬度、城市、省份和区县信息。客户端将 city、province、district、longitude、latitude、source、locatedAt 提交给服务端，由服务端写入 app-state.json 文件。"
    )
    add_body(
        document,
        "天气阶段由服务端读取最近一次定位信息后，先调用和风天气城市搜索接口获取标准化城市与 LocationID，再调用实时天气接口获取温度、体感温度、风向风速、湿度、能见度等数据，调用空气质量接口获取 AQI、分类、首要污染物和 PM2.5、PM10、NO2、SO2、CO、O3 等污染物浓度，并调用 24 小时逐小时天气接口生成温湿度折线图。"
    )
    add_bullet(document, "百度地图：用于定位与 IP 城市兜底。")
    add_bullet(document, "和风天气 GeoAPI：用于城市标准化和 LocationID 获取。")
    add_bullet(document, "和风天气实时天气 API：用于天气详情。")
    add_bullet(document, "和风天气空气质量 API v1：用于 AQI 与污染物浓度。")
    add_bullet(document, "和风天气 24 小时天气预报 API：用于温湿度折线图。")

    add_heading(document, "四、结果展示及分析")
    add_body(
        document,
        "系统首页可以完成定位城市的获取、显示与服务端保存，并由 Flask 服务端生成天气和空气质量数据，以卡片形式展示实时天气、AQI 和主要污染物，同时使用 SVG 折线图展示未来 24 小时的温度和湿度变化趋势。"
    )
    add_body(
        document,
        "从实现结果看，题目要求中的三项核心功能已经全部覆盖：一是定位结果被服务端保存并在客户端回显；二是移动端界面具有天气与空气质量的动态展示；三是天气详情和 AQI 由服务端根据定位城市统一获取并保存后再返回给首页 Body 区。该方案结构简单、实现成本低、适合课程测试环境。"
    )
    add_screenshot(document, "4.1 手机端单页截图", "home.png")

    add_heading(document, "五、测试与验收")
    add_score_table(document)
    add_bullet(document, "测试 1：点击定位按钮后，客户端显示城市、省份、经纬度和定位方式。")
    add_bullet(document, "测试 2：检查 data/app-state.json，确认最新定位结果已保存到服务端。")
    add_bullet(document, "测试 3：首页刷新后，确认天气、AQI、污染物和折线图成功展示。")
    add_bullet(document, "测试 4：重启服务端后再次访问，已保存的定位和天气快照文件仍然存在。")
    add_bullet(document, "测试 5：live 模式缺少和风天气环境变量时，服务端返回明确错误提示。")
    add_bullet(document, "测试 6：接口返回异常时，前端页面应给出错误提示，不出现空白页。")
    add_bullet(document, "自动化测试：uv run pytest 当前覆盖 Flask 路由、定位保存、mock dashboard、live 错误处理和持久化读写。")

    add_heading(document, "六、心得与体会")
    add_body(
        document,
        "本实验的重点不是堆叠功能，而是围绕题目要求构建一条完整的数据链路：定位、服务端保存、服务端查询天气、前端动态展示和实验报告整理。通过本次实验，可以进一步理解移动端页面适配、第三方接口调用、前后端数据传递和服务端持久化的基本方法。"
    )
    add_body(
        document,
        "在实现过程中，为了避免超出题目范围，系统没有增加历史查询、多账户管理或消息通知等附加功能，而是把精力集中在题目评分点对应的能力上。这种面向评分项的实现方式有利于在有限时间内稳定完成实验任务。"
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "附录：源码与运行说明")
    add_bullet(document, "启动命令：uv sync && uv run python app.py")
    add_bullet(document, "环境变量：HOST、PORT、APP_DATA_MODE、BAIDU_MAP_AK、QWEATHER_API_HOST、QWEATHER_API_KEY")
    add_bullet(document, "服务端持久化文件：data/app-state.json")
    add_bullet(document, "报告生成脚本：scripts/generate_report.py")
    add_bullet(document, "提交前事项：填写报告封面个人信息；真实演示时补齐百度地图 AK、和风天气 API Host 与 API Key。")

    document.save(str(OUTPUT_FILE))
    print(f"Report generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
